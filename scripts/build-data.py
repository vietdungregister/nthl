#!/usr/bin/env python3
"""
build_data.py — NTHL Data Pipeline
====================================
Chạy toàn bộ PHASE DATA.1 → DATA.9:
  1. Đọc dữ liệu Facebook JSON
  2. Dedup (theo timestamp + title)
  3. Lọc bỏ "shared a memory" → log works_removed_shares.json
  4. Phân loại tác phẩm (rule-based)
  5. Tạo works.json, poems_only.json, notes.json
  6. Chunking → chunks.json
  7. Media catalog → media_catalog.json
  8. Cập nhật file Word (loại share + tạo NTHL_Tho.docx)

Cách chạy:
  python3 build_data.py
"""

import json, re, uuid, datetime, os, unicodedata
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

# ── Paths ──────────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent
FB_DIR   = BASE_DIR / "facebook-nguyenthehoanglinh-06_03_2026-y1cYzpJi/your_facebook_activity"
POSTS_DIR = FB_DIR / "posts"
NOTES_FILE = FB_DIR / "other_activity/notes.json"
COMMENTS_FILE = FB_DIR / "comments_and_reactions/comments.json"
VIDEOS_FILE   = FB_DIR / "posts/your_videos.json"
MEDIA_DIR = BASE_DIR / "media"
OUTPUT_DIR = BASE_DIR / "output"
DATA_DIR   = OUTPUT_DIR / "data"

OUTPUT_DIR.mkdir(exist_ok=True)
DATA_DIR.mkdir(exist_ok=True)

# ── Encoding helper ─────────────────────────────────────────────────────────
def decode_fb(text):
    if not text: return ""
    try: return text.encode('latin-1').decode('utf-8')
    except: return text

def clean_xml_text(text):
    """Loại bỏ ký tự không hợp lệ trong XML/Word."""
    if not text: return ""
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)

# ── 1. Load & Dedup posts ───────────────────────────────────────────────────
def load_all_posts():
    print("📂 Đọc dữ liệu bài đăng...")
    all_posts = []
    for i in range(1, 8):
        path = POSTS_DIR / f"your_posts__check_ins__photos_and_videos_{i}.json"
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        all_posts.extend(data)
        print(f"   File {i}: {len(data):,} bài")
    all_posts.sort(key=lambda x: x.get('timestamp', 0))
    print(f"   Tổng: {len(all_posts):,} bài\n")
    return all_posts

def dedup_posts(all_posts):
    seen = {}
    for post in all_posts:
        key = (post.get('timestamp', 0), post.get('title', ''))
        text_len = sum(len(d.get('post', '')) for d in post.get('data', []) if 'post' in d)
        if key not in seen or text_len > seen[key][1]:
            seen[key] = (post, text_len)
    deduped = sorted([v[0] for v in seen.values()], key=lambda x: x.get('timestamp', 0))
    print(f"🧹 Dedup: {len(all_posts):,} → {len(deduped):,} bài (loại {len(all_posts)-len(deduped):,})\n")
    return deduped

# ── 2. Filter shared memories ────────────────────────────────────────────────
SHARE_TITLE_PATTERNS = [
    'shared a memory', 'shared a post', 'shared a photo',
    'shared a video', 'shared a link',
    'đã chia sẻ kỷ niệm', 'đã chia sẻ bài viết',
]

def is_shared_memory(post):
    title = decode_fb(post.get('title', '')).lower()
    return any(p in title for p in SHARE_TITLE_PATTERNS)

def filter_shares(all_posts):
    print("✂️  Lọc bài 'shared a memory'...")
    originals = []
    removed = []
    for post in all_posts:
        if is_shared_memory(post):
            ts = post.get('timestamp', 0)
            dt = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d') if ts else ''
            text = ''
            for d in post.get('data', []):
                if 'post' in d:
                    text = decode_fb(d['post'])[:200]
                    break
            removed.append({
                'timestamp': ts,
                'date': dt,
                'title': decode_fb(post.get('title', '')),
                'text_preview': text,
            })
        else:
            originals.append(post)

    # Log bài đã loại
    log_path = DATA_DIR / 'works_removed_shares.json'
    with open(log_path, 'w', encoding='utf-8') as f:
        json.dump(removed, f, ensure_ascii=False, indent=2)
    print(f"   Giữ lại: {len(originals):,} bài")
    print(f"   Đã loại: {len(removed):,} bài → {log_path.relative_to(BASE_DIR)}\n")
    return originals

# ── 3. Classify works ────────────────────────────────────────────────────────
def get_media_type(post):
    """Trả về ('video'|'image'|None)."""
    for att in post.get('attachments', []):
        for item in att.get('data', []):
            if 'media' in item:
                media = item['media']
                uri = media.get('uri', '')
                if uri.endswith('.mp4') or 'videos' in uri:
                    return 'video'
                return 'image'
    return None

def classify_work(text, title, media_type):
    lines = [l for l in text.split('\n') if l.strip()]
    if not lines:
        if media_type == 'video': return 'video'
        if media_type == 'image': return 'photo'
        return 'status'

    avg_line_len = sum(len(l) for l in lines) / len(lines)
    title_lower = title.lower()

    if media_type == 'video':
        if 'fifa' in text.lower() or 'fifa' in title_lower:
            return 'video_fifa'
        if any(k in text.lower() or k in title_lower for k in ['chatgpt', 'chat gpt', ' gpt', 'openai']):
            return 'video_chatgpt'
        return 'video'

    if media_type == 'image' and len(text) < 50:
        return 'photo'

    # Thơ: ≥4 dòng, dòng ngắn trung bình <50 ký tự
    if len(lines) >= 4 and avg_line_len < 50:
        return 'poem'

    # Tiểu luận/truyện: văn xuôi dài, ít dòng
    if len(text) > 300 and len(lines) <= max(3, len(text) // 80):
        return 'prose'

    return 'status'

# ── 4. Load comments ─────────────────────────────────────────────────────────
def load_own_comments():
    with open(COMMENTS_FILE, 'r', encoding='utf-8') as f:
        data = json.load(f)
    comments = []
    for item in data.get('comments_v2', []):
        title = decode_fb(item.get('title', ''))
        if 'commented on his own' not in title and 'own post' not in title.lower():
            continue
        for d in item.get('data', []):
            if 'comment' in d:
                c = d['comment']
                ts = c.get('timestamp', item.get('timestamp', 0))
                text = decode_fb(c.get('comment', ''))
                if text:
                    comments.append({'ts': ts, 'text': text})
    comments.sort(key=lambda x: x['ts'])
    print(f"💬 Comments: {len(comments):,}\n")
    return comments

# ── 5. Build works data ──────────────────────────────────────────────────────
def auto_title(text, dt_str):
    """Dòng đầu tiên có ký tự, tối đa 60 ký tự. Nếu rỗng → dùng ngày."""
    for line in text.split('\n'):
        line = line.strip()
        if len(line) >= 3:
            return line[:60] + ('...' if len(line) > 60 else '')
    return f"Bài viết ngày {dt_str}"

def vn_slugify(text, max_length=80):
    """Tạo slug từ tiếng Việt không cần thư viện ngoài."""
    # Normalize unicode → decompose accents
    text = unicodedata.normalize('NFD', text)
    text = ''.join(c for c in text if unicodedata.category(c) != 'Mn')
    # Lowercase, replace spaces/special chars with hyphens
    text = text.lower()
    text = re.sub(r'[^a-z0-9\s-]', '', text)
    text = re.sub(r'[\s-]+', '-', text).strip('-')
    return text[:max_length]

def auto_slug(title, ts):
    base = vn_slugify(title)
    if not base:
        base = f"bai-viet-{ts}"
    return base

def make_unique_slug(slug, seen_slugs):
    original = slug
    counter = 2
    while slug in seen_slugs:
        slug = f"{original}-{counter}"
        counter += 1
    seen_slugs.add(slug)
    return slug

def build_works(posts):
    print("🔨 Xây dựng works data...")
    works = []
    seen_slugs = set()
    genre_counts = {}
    video_ids = set()

    # Load video descriptions for matching
    try:
        with open(VIDEOS_FILE, 'r', encoding='utf-8') as f:
            vid_data = json.load(f)
        for v in vid_data.get('videos_v2', []):
            desc = decode_fb(v.get('description', '')).lower()
            video_ids.add(desc[:50])  # rough match
    except: pass

    for post in posts:
        ts = post.get('timestamp', 0)
        dt = datetime.datetime.fromtimestamp(ts) if ts else datetime.datetime(2009, 1, 1)
        dt_str = dt.strftime('%Y-%m-%d')
        year = dt.year

        text = ''
        for d in post.get('data', []):
            if 'post' in d:
                text = decode_fb(d['post'])
                break

        title_raw = decode_fb(post.get('title', ''))
        media_type = get_media_type(post)
        genre = classify_work(text, title_raw, media_type)
        genre_counts[genre] = genre_counts.get(genre, 0) + 1

        # DB genre mapping
        db_genre_map = {
            'poem': 'poem', 'prose': 'prose', 'status': 'prose',
            'video': 'video', 'video_fifa': 'video', 'video_chatgpt': 'video',
            'photo': 'photo'
        }
        db_genre = db_genre_map.get(genre, 'prose')

        # Sub-genre tag for video
        sub_genre = None
        if genre == 'video_fifa': sub_genre = 'video-fifa'
        elif genre == 'video_chatgpt': sub_genre = 'video-chatgpt'

        title = auto_title(text, dt_str) if text else (title_raw or f"Bài viết {dt_str}")
        slug_base = auto_slug(title, ts)
        slug = make_unique_slug(slug_base, seen_slugs)

        # Excerpt: 150 ký tự đầu
        excerpt = (text[:150] + '...') if len(text) > 150 else text

        tags = [f"auto:{year}", f"auto:{db_genre}"]
        if sub_genre: tags.append(f"auto:{sub_genre}")

        comments = []  # sẽ match sau

        work = {
            'id': str(uuid.uuid4()),
            'title': title,
            'slug': slug,
            'genre': db_genre,
            'autoGenre': genre,  # raw classified genre
            'content': text,
            'excerpt': excerpt,
            'status': 'published',
            'publishedAt': dt.isoformat() + 'Z',
            'createdAt': dt.isoformat() + 'Z',
            'tags': tags,
            'source': 'facebook',
            'fbTimestamp': ts,
            'autoClassified': True,
            'comments': comments,
        }
        works.append(work)

    print(f"   Tổng works: {len(works):,}")
    print(f"   Phân loại:")
    for g, c in sorted(genre_counts.items(), key=lambda x: -x[1]):
        print(f"     {g}: {c:,}")
    print()
    return works

# ── 6. Notes ─────────────────────────────────────────────────────────────────
def clean_html_note(text):
    if not text: return ''
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<p[^>]*>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</p>', '', text, flags=re.IGNORECASE)
    text = re.sub(r'<[^>]+>', '', text)
    text = text.replace('\xa0', ' ')
    text = re.sub(r'[ \t]{3,}', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def load_notes():
    print("📝 Đọc Notes...")
    with open(NOTES_FILE, 'r', encoding='utf-8') as f:
        data = json.load(f)
    notes = []
    for item in data.get('notes_v2', []):
        ts = item.get('timestamp', 0)
        dt = datetime.datetime.fromtimestamp(ts).isoformat() + 'Z' if ts else None
        title = decode_fb(item.get('title', ''))
        content_raw = ''
        for d in item.get('data', []):
            if isinstance(d, dict) and 'note' in d:
                content_raw = decode_fb(d.get('note', ''))
        content = clean_html_note(content_raw)
        notes.append({
            'id': str(uuid.uuid4()),
            'title': title or f"Ghi chú {dt[:10] if dt else ''}",
            'content': content,
            'publishedAt': dt,
            'source': 'facebook_notes',
        })
    print(f"   Tổng: {len(notes):,} notes\n")
    return notes

# ── 7. Chunking ───────────────────────────────────────────────────────────────
def chunk_work(work):
    """Chia bài thành chunks 1-2 dòng cho chatbot."""
    text = work.get('content', '')
    if not text:
        return []
    chunks = []
    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if len(line) < 8:
            i += 1
            continue
        # Dòng dài → chia theo câu
        if len(line) > 120:
            for s in re.split(r'(?<=[.!?])\s+', line):
                s = s.strip()
                if len(s) >= 8:
                    chunks.append({'content': s})
            i += 1
            continue
        # Ghép 2 dòng ngắn
        if len(line) < 60 and i + 1 < len(lines):
            next_line = lines[i+1].strip()
            if len(next_line) >= 8 and len(next_line) < 60:
                chunks.append({'content': f"{line}\n{next_line}"})
                i += 2
                continue
        chunks.append({'content': line})
        i += 1
    return chunks

def build_chunks(works):
    print("✂️  Chunking works...")
    all_chunks = []
    for work in works:
        if work['genre'] in ('video', 'photo'):
            continue  # video/photo không chunk text
        raw_chunks = chunk_work(work)
        for c in raw_chunks:
            all_chunks.append({
                'id': str(uuid.uuid4()),
                'workId': work['id'],
                'workTitle': work['title'],
                'workSlug': work['slug'],
                'content': c['content'],
                'score': 0,
                'isBlocked': False,
            })
    print(f"   Tổng chunks: {len(all_chunks):,}\n")
    return all_chunks

# ── 8. Media catalog ──────────────────────────────────────────────────────────
def build_media_catalog():
    print("📷 Xây dựng media catalog...")
    catalog = []

    # Load video descriptions for category matching
    video_descriptions = {}
    try:
        with open(VIDEOS_FILE, 'r', encoding='utf-8') as f:
            vid_data = json.load(f)
        for v in vid_data.get('videos_v2', []):
            uri = v.get('uri', '')
            desc = decode_fb(v.get('description', ''))
            if uri:
                filename = Path(uri).name
                video_descriptions[filename] = desc
    except Exception as e:
        print(f"   ⚠️ Không đọc được videos.json: {e}")

    # Videos
    videos_dir = MEDIA_DIR / 'videos'
    if videos_dir.exists():
        for f in sorted(videos_dir.glob('*.mp4')):
            desc = video_descriptions.get(f.name, '')
            desc_lower = desc.lower()
            if 'fifa' in desc_lower:
                category = 'video_fifa'
            elif any(k in desc_lower for k in ['chatgpt', 'chat gpt', 'gpt', 'openai']):
                category = 'video_chatgpt'
            else:
                category = 'video'

            stat = f.stat()
            # Try to get date from FB metadata
            ts = None
            try:
                ts = int(stat.st_birthtime)
            except:
                ts = int(stat.st_mtime)

            catalog.append({
                'id': str(uuid.uuid4()),
                'filename': f.name,
                'path': str(f.relative_to(BASE_DIR)),
                'type': 'video',
                'category': category,
                'size': stat.st_size,
                'description': desc[:200] if desc else '',
                'linkedWorkId': None,
            })

    video_count = len([c for c in catalog if c['type'] == 'video'])

    # Images
    img_extensions = {'.jpg', '.jpeg', '.png', '.webp', '.gif'}
    for album_dir in sorted(MEDIA_DIR.iterdir()):
        if not album_dir.is_dir() or album_dir.name == 'videos':
            continue
        album_name = album_dir.name
        for f in sorted(album_dir.iterdir()):
            if f.suffix.lower() not in img_extensions:
                continue
            stat = f.stat()
            catalog.append({
                'id': str(uuid.uuid4()),
                'filename': f.name,
                'path': str(f.relative_to(BASE_DIR)),
                'type': 'image',
                'category': 'photo',  # tác giả tự tag tranh/ảnh sau
                'album': album_name,
                'size': stat.st_size,
                'linkedWorkId': None,
            })

    image_count = len([c for c in catalog if c['type'] == 'image'])
    print(f"   Videos: {video_count:,} | Images: {image_count:,} | Tổng: {len(catalog):,}\n")
    return catalog

# ── 9. Word files ─────────────────────────────────────────────────────────────
def make_docx_paragraph(doc, text, style='Normal', font_name='Times New Roman', font_size=12):
    """Thêm paragraph vào doc, xử lý line breaks."""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(0)
    lines = text.split('\n')
    for j, line in enumerate(lines):
        if j > 0:
            br = OxmlElement('w:br')
            p._p.append(br)
        run = p.add_run(clean_xml_text(line))
        run.font.name = font_name
        run.font.size = Pt(font_size)
        r = run._element
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rPr.append(rFonts)
    return p

def save_works_to_docx(works_by_year, output_dir):
    """Tạo 1 file Word/năm cho tất cả bài đã lọc."""
    print("📄 Tạo file Word (đã loại share)...")
    for label, year_works in sorted(works_by_year.items()):
        doc = Document()
        doc.core_properties.title = f"NTHL {label}"
        for style in doc.styles:
            if style.name == 'Normal':
                style.font.name = 'Times New Roman'
                style.font.size = Pt(12)

        for work in year_works:
            ts = work['fbTimestamp']
            dt_str = datetime.datetime.fromtimestamp(ts).strftime('%d/%m/%Y') if ts else ''
            title = work['title']
            content = work['content']
            genre_label = work['autoGenre']

            # Date header
            make_docx_paragraph(doc, f"[{dt_str}] {title}", font_size=11)

            # Genre badge nếu video/photo
            if genre_label in ('video', 'video_fifa', 'video_chatgpt', 'photo'):
                badges = {'video': '🎬 Video', 'video_fifa': '⚽ Video Fifa',
                          'video_chatgpt': '🤖 Video ChatGPT', 'photo': '📷 Ảnh'}
                make_docx_paragraph(doc, f"[{badges.get(genre_label, genre_label)}]", font_size=10)

            if content:
                make_docx_paragraph(doc, content)

            # Comments
            for cm in work.get('comments', []):
                make_docx_paragraph(doc, f"    💬 {cm['text']}", font_size=10)

            doc.add_paragraph()  # blank line

        out_path = output_dir / f"NTHL_{label}.docx"
        doc.save(out_path)
        print(f"   ✅ NTHL_{label}.docx ({len(year_works):,} bài)")

def save_poems_docx(poem_works, output_dir):
    """Tạo NTHL_Tho.docx — thư viện thơ riêng."""
    print("\n📚 Tạo NTHL_Tho.docx...")
    doc = Document()
    doc.core_properties.title = "NTHL — Thư Viện Thơ"
    for style in doc.styles:
        if style.name == 'Normal':
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)

    poem_works_sorted = sorted(poem_works, key=lambda w: w['fbTimestamp'])
    for work in poem_works_sorted:
        ts = work['fbTimestamp']
        dt_str = datetime.datetime.fromtimestamp(ts).strftime('%d/%m/%Y') if ts else ''
        make_docx_paragraph(doc, f"[{dt_str}]", font_size=10)
        if work['content']:
            make_docx_paragraph(doc, work['content'])
        doc.add_paragraph()

    out_path = output_dir / "NTHL_Tho.docx"
    doc.save(out_path)
    print(f"   ✅ NTHL_Tho.docx ({len(poem_works_sorted):,} bài thơ)\n")

# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    print("=" * 60)
    print("  NTHL Data Pipeline — build_data.py")
    print("=" * 60)
    print()

    # Step 1-2: Load, dedup, filter
    all_posts = load_all_posts()
    all_posts = dedup_posts(all_posts)
    original_posts = filter_shares(all_posts)

    # Step 3: Load comments (for future enrichment)
    own_comments = load_own_comments()

    # Step 4: Build works with classification
    works = build_works(original_posts)

    # Step 5: Load and save Notes
    notes = load_notes()

    # Step 6: Chunking
    chunks = build_chunks(works)

    # Step 7: Media catalog
    media_catalog = build_media_catalog()

    # ── Save JSON outputs ──
    print("💾 Lưu JSON...")

    works_path = DATA_DIR / 'works.json'
    with open(works_path, 'w', encoding='utf-8') as f:
        json.dump(works, f, ensure_ascii=False, indent=2)
    print(f"   ✅ works.json ({len(works):,} bài)")

    poems = [w for w in works if w['genre'] == 'poem']
    poems_path = DATA_DIR / 'poems_only.json'
    with open(poems_path, 'w', encoding='utf-8') as f:
        json.dump(poems, f, ensure_ascii=False, indent=2)
    print(f"   ✅ poems_only.json ({len(poems):,} bài thơ)")

    notes_path = DATA_DIR / 'notes.json'
    with open(notes_path, 'w', encoding='utf-8') as f:
        json.dump(notes, f, ensure_ascii=False, indent=2)
    print(f"   ✅ notes.json ({len(notes):,} ghi chú)")

    chunks_path = DATA_DIR / 'chunks.json'
    with open(chunks_path, 'w', encoding='utf-8') as f:
        json.dump(chunks, f, ensure_ascii=False, indent=2)
    print(f"   ✅ chunks.json ({len(chunks):,} chunks)")

    media_path = DATA_DIR / 'media_catalog.json'
    with open(media_path, 'w', encoding='utf-8') as f:
        json.dump(media_catalog, f, ensure_ascii=False, indent=2)
    print(f"   ✅ media_catalog.json ({len(media_catalog):,} files)")
    print()

    # ── Save Word files ──
    # Group by year for posts
    works_by_year = {}
    for w in works:
        ts = w['fbTimestamp']
        year = datetime.datetime.fromtimestamp(ts).year if ts else 0
        if year < 2012:
            label = '2009-2012'
        else:
            label = str(year)
        works_by_year.setdefault(label, []).append(w)

    save_works_to_docx(works_by_year, OUTPUT_DIR)
    save_poems_docx(poems, OUTPUT_DIR)

    # ── Summary ──
    print()
    print("=" * 60)
    print("  ✅ HOÀN THÀNH!")
    print(f"     - {len(works):,} bài gốc (đã lọc share)")
    print(f"     - {len(poems):,} bài thơ → NTHL_Tho.docx")
    print(f"     - {len(notes):,} notes → notes.json")
    print(f"     - {len(chunks):,} chunks → chunks.json")
    print(f"     - {len(media_catalog):,} media files → media_catalog.json")
    print(f"     - {len([c for c in media_catalog if c['category']=='video_fifa'])} video Fifa")
    print(f"     - {len([c for c in media_catalog if c['category']=='video_chatgpt'])} video ChatGPT")
    print(f"")
    print(f"  📁 Output: {OUTPUT_DIR}/")
    print(f"  📁 Data:   {DATA_DIR}/")
    print("=" * 60)

if __name__ == '__main__':
    main()
