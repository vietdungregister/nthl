"""
export_to_docx.py
Chuyển đổi dữ liệu Facebook của Nguyễn Thế Hoàng Linh → file Word (.docx)
Đầu ra: 15 file posts (theo năm) + 1 file NTHL_Notes.docx tách riêng
"""

import json
import bisect
import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
# PATHS
# ─────────────────────────────────────────────
BASE_DIR = Path("facebook-nguyenthehoanglinh-06_03_2026-y1cYzpJi")
POSTS_DIR = BASE_DIR / "your_facebook_activity" / "posts"
COMMENTS_FILE = BASE_DIR / "your_facebook_activity" / "comments_and_reactions" / "comments.json"
NOTES_FILE = BASE_DIR / "your_facebook_activity" / "other_activity" / "notes.json"
OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)


# ─────────────────────────────────────────────
# TEXT HELPERS
# ─────────────────────────────────────────────

def decode_fb_text(text):
    """Decode Facebook's Latin-1 encoded UTF-8 text."""
    if not text:
        return ""
    try:
        return text.encode('latin-1').decode('utf-8')
    except (UnicodeDecodeError, UnicodeEncodeError):
        return text


def clean_xml_text(text):
    """Remove NULL bytes and XML-invalid control characters."""
    if not text:
        return ""
    return ''.join(
        ch for ch in text
        if ch in ('\t', '\n', '\r')
        or (0x20 <= ord(ch) <= 0xD7FF)
        or (0xE000 <= ord(ch) <= 0xFFFD)
        or (0x10000 <= ord(ch) <= 0x10FFFF)
    )


def add_text_with_linebreaks(paragraph, text, font_size=Pt(12), color=None, bold=False, italic=False):
    """Add text to paragraph preserving line breaks (critical for poetry)."""
    text = clean_xml_text(text)
    lines = text.split('\n')
    for i, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.font.name = 'Times New Roman'
        run.font.size = font_size
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        if i < len(lines) - 1:
            br = OxmlElement('w:br')
            run._r.append(br)


# ─────────────────────────────────────────────
# POST CLASSIFICATION
# ─────────────────────────────────────────────

def classify_post(post):
    """
    Phân loại bài đăng:
      'shared_memory' — share lại kỷ niệm từ năm trước
      'shared_other'  — share bài/link/reel/ảnh của người khác
      'original'      — bài gốc
    """
    title = decode_fb_text(post.get('title', ''))
    if 'shared a memory' in title:
        return 'shared_memory'
    share_keywords = [
        'shared a post', 'shared a link', 'shared a note',
        'shared a reel', 'shared a photo', 'Shared from Instagram',
        'shared an event', 'shared a Page', 'shared an album',
        'shared an episode',
    ]
    for kw in share_keywords:
        if kw in title:
            return 'shared_other'
    return 'original'


def extract_memory_content(post):
    """
    Từ bài 'shared a memory', trích xuất ngày gốc + nội dung gốc.
    Cấu trúc attachments.data: [ago_text, old_title, original_date, original_content, ...]
    """
    if 'attachments' not in post:
        return None, None
    for att in post['attachments']:
        texts = [
            decode_fb_text(d.get('text', ''))
            for d in att.get('data', [])
            if 'text' in d
        ]
        if len(texts) >= 4:
            return texts[2], texts[3]   # original_date_str, original_text
        elif len(texts) == 3:
            return texts[1], texts[2]
        elif len(texts) == 2:
            return None, texts[1]
        elif len(texts) == 1:
            return None, texts[0]
    return None, None


def parse_fb_date_str(date_str):
    """Parse chuỗi ngày FB "Dec 01, 2018 11:55:49 pm" → datetime hoặc None."""
    if not date_str:
        return None
    for fmt in ["%b %d, %Y %I:%M:%S %p", "%b %d, %Y"]:
        try:
            return datetime.datetime.strptime(date_str.strip(), fmt)
        except ValueError:
            continue
    return None


def get_post_content(post):
    """
    Trả về (post_type, text, original_date_str).
      post_type: 'original' / 'shared_memory' / 'shared_other'
      text: nội dung đã decode
      original_date_str: ngày gốc (chỉ với shared_memory)
    """
    ptype = classify_post(post)

    if ptype == 'shared_memory':
        orig_date_str, orig_text = extract_memory_content(post)
        text = orig_text if orig_text else ""
        # Fallback sang data.post nếu attachment không có text
        if not text or text == '.':
            for d in post.get('data', []):
                if 'post' in d:
                    t = decode_fb_text(d['post'])
                    if t and t != '.':
                        text = t
                        break
        return ptype, text, orig_date_str

    # original hoặc shared_other
    text = ""
    for d in post.get('data', []):
        if 'post' in d:
            t = decode_fb_text(d['post'])
            if t and t != '.':
                text = t
                break

    if not text:
        if 'attachments' in post:
            for att in post['attachments']:
                for ad in att.get('data', []):
                    if 'media' in ad:
                        text = "[Ảnh/Video]"
                        break
                    if 'external_context' in ad:
                        url = ad['external_context'].get('url', '')
                        text = f"[Liên kết: {url[:80]}]" if url else "[Liên kết]"
                        break
                if text:
                    break
        if not text:
            text = "[Không có nội dung]"

    return ptype, text, None


# ─────────────────────────────────────────────
# LOAD DATA
# ─────────────────────────────────────────────

def load_all_posts():
    print("📂 Đang đọc dữ liệu bài đăng...")
    all_posts = []
    for i in range(1, 8):
        path = POSTS_DIR / f"your_posts__check_ins__photos_and_videos_{i}.json"
        print(f"   Đọc file {i}/7: {path.name}", end="", flush=True)
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        all_posts.extend(data)
        print(f" → {len(data):,} bài")
    all_posts.sort(key=lambda x: x.get('timestamp', 0))
    print(f"   ✅ Tổng: {len(all_posts):,} bài đăng\n")
    return all_posts


def dedup_posts(all_posts):
    """
    Loại bỏ bài trùng lặp do Facebook export bug.
    Key: (timestamp, title) — giữ bản có text dài nhất.
    """
    seen = {}  # key → (post, text_len)
    for post in all_posts:
        ts = post.get('timestamp', 0)
        title = post.get('title', '')
        key = (ts, title)
        text_len = 0
        for d in post.get('data', []):
            if 'post' in d:
                text_len = max(text_len, len(d.get('post', '')))
        if key not in seen or text_len > seen[key][1]:
            seen[key] = (post, text_len)

    deduped = [v[0] for v in seen.values()]
    deduped.sort(key=lambda x: x.get('timestamp', 0))
    removed = len(all_posts) - len(deduped)
    print(f"🧹 Dedup: {len(all_posts):,} → {len(deduped):,} bài (loại bỏ {removed:,} bản trùng)\n")
    return deduped


def load_notes():
    print("📝 Đang đọc dữ liệu Notes...")
    with open(NOTES_FILE, 'r', encoding='utf-8') as f:
        data = json.load(f)
    notes = data.get('notes_v2', [])
    notes.sort(key=lambda x: x.get('created_timestamp', 0))
    print(f"   ✅ Tổng: {len(notes):,} ghi chú (Notes)\n")
    return notes


def load_own_comments():
    print("💬 Đang đọc dữ liệu comment...")
    with open(COMMENTS_FILE, 'r', encoding='utf-8') as f:
        data = json.load(f)
    own_comments = []
    for c in data.get('comments_v2', []):
        title = decode_fb_text(c.get('title', ''))
        if 'his own' not in title and 'her own' not in title:
            continue
        comment_text = ""
        for d in c.get('data', []):
            if isinstance(d, dict) and 'comment' in d and isinstance(d['comment'], dict):
                comment_text = decode_fb_text(d['comment'].get('comment', ''))
                break
        own_comments.append({
            'timestamp': c.get('timestamp', 0),
            'text': comment_text,
        })
    own_comments.sort(key=lambda x: x['timestamp'])
    print(f"   ✅ Tổng: {len(own_comments):,} comment trên bài của mình\n")
    return own_comments


def match_comments_to_posts(all_posts, own_comments):
    """Gắn comment vào bài đăng gần nhất trước đó (binary search)."""
    print("🔗 Đang ghép comment vào bài đăng...")
    post_timestamps = [p.get('timestamp', 0) for p in all_posts]
    for post in all_posts:
        post['_comments'] = []
    matched = 0
    for comment in own_comments:
        ct = comment['timestamp']
        idx = bisect.bisect_right(post_timestamps, ct) - 1
        if idx >= 0:
            all_posts[idx]['_comments'].append(comment)
            matched += 1
    print(f"   ✅ Đã ghép {matched:,}/{len(own_comments):,} comment\n")
    return all_posts


def group_posts_by_year(all_posts):
    """Nhóm bài theo năm. 2009-2012 gộp chung."""
    groups = {}
    for post in all_posts:
        year = datetime.datetime.fromtimestamp(post.get('timestamp', 0)).year
        label = "2009-2012" if year <= 2012 else str(year)
        groups.setdefault(label, []).append(post)
    return groups


# ─────────────────────────────────────────────
# DOCX HELPERS
# ─────────────────────────────────────────────

def set_document_style(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing = Pt(16)


def clear_default_paragraph(doc):
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)


def add_title_page(doc, title_text, subtitle_text):
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(30)
    run.bold = True
    run.font.color.rgb = RGBColor(30, 30, 30)
    title_p.paragraph_format.space_after = Pt(12)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run(subtitle_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)
    sub_p.paragraph_format.space_after = Pt(0)

    doc.add_page_break()


def add_separator(doc):
    p = doc.add_paragraph()
    run = p.add_run("─" * 50)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(180, 180, 180)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def format_date(timestamp):
    return datetime.datetime.fromtimestamp(timestamp).strftime('%d/%m/%Y')


def format_original_date(date_str):
    dt = parse_fb_date_str(date_str)
    return dt.strftime('%d/%m/%Y') if dt else (date_str or "")


# ─────────────────────────────────────────────
# GENERATE POSTS DOCX
# ─────────────────────────────────────────────

def generate_posts_docx(year_label, posts):
    doc = Document()
    set_document_style(doc)
    clear_default_paragraph(doc)
    add_title_page(doc, "Nguyễn Thế Hoàng Linh", f"Tuyển tập bài viết {year_label}")

    counts = {'original': 0, 'shared_memory': 0, 'shared_other': 0}

    for post in posts:
        ts = post.get('timestamp', 0)
        post_date = format_date(ts)
        ptype, text, orig_date_str = get_post_content(post)
        counts[ptype] = counts.get(ptype, 0) + 1

        # Separator
        add_separator(doc)

        # Ngày đăng
        date_p = doc.add_paragraph()
        date_p.paragraph_format.space_before = Pt(0)
        date_p.paragraph_format.space_after = Pt(2)
        run = date_p.add_run(f"📅 {post_date}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor(80, 80, 80)

        # Nhãn loại bài
        if ptype == 'shared_memory':
            orig_date_fmt = format_original_date(orig_date_str) if orig_date_str else "năm trước"
            share_p = doc.add_paragraph()
            share_p.paragraph_format.space_before = Pt(0)
            share_p.paragraph_format.space_after = Pt(4)
            run = share_p.add_run(f"🔄 Chia sẻ lại bài viết ngày {orig_date_fmt}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.italic = True
            run.font.color.rgb = RGBColor(120, 100, 60)

        elif ptype == 'shared_other':
            title_raw = decode_fb_text(post.get('title', ''))
            if 'Instagram' in title_raw:
                label = "📷 Từ Instagram"
            elif 'link' in title_raw:
                label = "🔗 Chia sẻ liên kết"
            elif 'note' in title_raw:
                label = "📝 Chia sẻ ghi chú"
            else:
                label = "🔗 Chia sẻ"
            share_p = doc.add_paragraph()
            share_p.paragraph_format.space_before = Pt(0)
            share_p.paragraph_format.space_after = Pt(4)
            run = share_p.add_run(label)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.italic = True
            run.font.color.rgb = RGBColor(80, 100, 130)

        # Nội dung
        if text and text not in ('.', ''):
            content_p = doc.add_paragraph()
            content_p.paragraph_format.space_before = Pt(4)
            content_p.paragraph_format.space_after = Pt(6)
            add_text_with_linebreaks(content_p, text, font_size=Pt(12))

        # Comments
        comments = post.get('_comments', [])
        if comments:
            ch_p = doc.add_paragraph()
            ch_p.paragraph_format.left_indent = Inches(0.4)
            ch_p.paragraph_format.space_before = Pt(4)
            ch_p.paragraph_format.space_after = Pt(2)
            run = ch_p.add_run("💬 Bình luận của tác giả:")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)

            for comment in comments:
                comment_text = comment.get('text', '').strip()
                if not comment_text:
                    continue
                cp = doc.add_paragraph()
                cp.paragraph_format.left_indent = Inches(0.4)
                cp.paragraph_format.space_before = Pt(0)
                cp.paragraph_format.space_after = Pt(4)
                add_text_with_linebreaks(cp, comment_text, font_size=Pt(11),
                                         color=RGBColor(60, 60, 60))

    out_path = OUTPUT_DIR / f"NTHL_{year_label}.docx"
    doc.save(str(out_path))
    return counts


# ─────────────────────────────────────────────
# GENERATE NOTES DOCX (tách riêng)
# ─────────────────────────────────────────────

def generate_notes_docx(notes):
    print("   Đang tạo NTHL_Notes.docx...", end="", flush=True)
    doc = Document()
    set_document_style(doc)
    clear_default_paragraph(doc)
    add_title_page(doc, "Nguyễn Thế Hoàng Linh", "Tuyển tập Ghi chú (Notes)")

    for note in notes:
        ts = note.get('created_timestamp', 0)
        note_date = format_date(ts)
        title = decode_fb_text(note.get('title', ''))
        text = decode_fb_text(note.get('text', ''))

        # Separator
        add_separator(doc)

        # Ngày viết
        date_p = doc.add_paragraph()
        date_p.paragraph_format.space_before = Pt(0)
        date_p.paragraph_format.space_after = Pt(2)
        run = date_p.add_run(f"📅 {note_date}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor(80, 80, 80)

        # Tiêu đề note
        if title:
            title_p = doc.add_paragraph()
            title_p.paragraph_format.space_before = Pt(0)
            title_p.paragraph_format.space_after = Pt(6)
            run = title_p.add_run(f"📝 {title}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(40, 90, 40)

        # Nội dung
        if text and text.strip():
            content_p = doc.add_paragraph()
            content_p.paragraph_format.space_before = Pt(4)
            content_p.paragraph_format.space_after = Pt(8)
            add_text_with_linebreaks(content_p, text, font_size=Pt(12))

    out_path = OUTPUT_DIR / "NTHL_Notes.docx"
    doc.save(str(out_path))
    print(f" ✅ → output/NTHL_Notes.docx ({len(notes):,} ghi chú)")


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────

def main():
    print("=" * 60)
    print("  Export dữ liệu Facebook → Word (.docx)")
    print("  Nguyễn Thế Hoàng Linh")
    print("=" * 60)
    print()

    # Load
    all_posts = load_all_posts()
    all_posts = dedup_posts(all_posts)
    notes = load_notes()
    own_comments = load_own_comments()

    # Match comments → posts
    all_posts = match_comments_to_posts(all_posts, own_comments)

    # Group posts by year
    groups = group_posts_by_year(all_posts)
    sorted_labels = sorted(groups.keys(), key=lambda x: int(x.split('-')[0]))

    # ── Generate 15 posts files ──
    print(f"📄 Bắt đầu tạo {len(sorted_labels)} file posts...\n")
    grand_total = 0
    for label in sorted_labels:
        posts = groups[label]
        print(f"   Đang xử lý {label} ({len(posts):,} bài)...", end="", flush=True)
        counts = generate_posts_docx(label, posts)
        grand_total += len(posts)
        print(f" ✅ → output/NTHL_{label}.docx")
        print(f"      (Bài gốc: {counts['original']:,} | "
              f"Chia sẻ kỷ niệm: {counts['shared_memory']:,} | "
              f"Chia sẻ khác: {counts['shared_other']:,})")

    # ── Generate Notes file ──
    print(f"\n📝 Tạo file Notes tách riêng...")
    generate_notes_docx(notes)

    print()
    print("=" * 60)
    print(f"  ✅ HOÀN THÀNH! Đã tạo {len(sorted_labels) + 1} file trong output/")
    print(f"     - {len(sorted_labels)} file bài đăng ({grand_total:,} bài)")
    print(f"     - 1 file NTHL_Notes.docx ({len(notes):,} ghi chú)")
    print("=" * 60)


if __name__ == "__main__":
    main()
